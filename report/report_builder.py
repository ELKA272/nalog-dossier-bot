import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from config import BASE_DIR, REPORTS_DIR


class ReportBuilder:
    def __init__(self, company_name: str, inn: str):
        self.company_name = company_name or "Организация"
        self.inn = inn
        self.doc = Document()

    def build(self, data: dict) -> str:
        self._add_title()
        self._add_section("1. Резюме для собственника бизнеса", self._resume(data))
        self._add_section("2. Общая информация", self._general(data))
        self._add_section("3. Руководители и учредители", self._management(data))
        self._add_section("4. Связанные компании", self._related(data))
        self._add_section("5. Финансовый анализ", self._finance(data))
        self._add_section("6. Судебные дела", self._court(data))
        self._add_section("7. Исполнительные производства", self._enforcement(data))
        self._add_section("8. Лицензии", self._licenses(data))
        self._add_risks_table(data)
        self._add_section("10. Вероятные претензии ФНС", self._fns_claims(data))
        self._add_section("11. Рекомендации по снижению рисков", self._recommendations(data))
        self._add_section("12. Использованные источники", self._sources(data))
        return self._save()

    def _add_title(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("НАЛОГОВОЕ ДОСЬЕ")
        run.bold = True
        run.font.size = Pt(18)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(
            f"{self.company_name}\nИНН: {self.inn}\n"
            f"Дата: {datetime.now().strftime('%d.%m.%Y')}"
        )
        run2.font.size = Pt(12)
        self.doc.add_paragraph("_" * 70)

    def _add_section(self, title, content_lines):
        if not content_lines:
            return
        p = self.doc.add_paragraph()
        run = p.add_run(f"\n{title}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)
        for line in content_lines:
            if line.strip():
                self.doc.add_paragraph(line.strip())

    def _get(self, d, *keys, default="—"):
        for k in keys:
            if isinstance(d, dict):
                d = d.get(k, {})
            else:
                return default
        if d is None or d == {}:
            return default
        return str(d) if d != "" else default

    def _fmt(self, val):
        if val is None or val == "" or val == 0:
            return "—"
        if isinstance(val, str):
            val = val.replace(" ", "")
            try:
                v = float(val)
            except ValueError:
                return val
        else:
            v = float(val)
        if abs(v) >= 1_000_000:
            return f"{v / 1_000_000:.2f} млн ₽"
        if abs(v) >= 1_000:
            return f"{v / 1_000:.1f} тыс ₽"
        return str(v)

    def _resume(self, data: dict) -> list:
        sc = data.get("scoring", {})
        lines = [
            f"Проведён анализ организации {self.company_name} (ИНН {self.inn}).",
            "",
            "Результаты оценки рисков:",
            f"  • Налоговый риск:      {sc.get('tax_score', 0)}/100 ({sc.get('tax_level', '—')})",
            f"  • Финансовый риск:     {sc.get('financial_score', 0)}/100 ({sc.get('financial_level', '—')})",
            f"  • Корпоративный риск:  {sc.get('corporate_score', 0)}/100 ({sc.get('corporate_level', '—')})",
            f"  • Репутационный риск:  {sc.get('reputation_score', 0)}/100 ({sc.get('reputation_level', '—')})",
            "",
            f"  • Риск дробления бизнеса:          {data.get('splitting', {}).get('conclusion', '—')}",
            f"  • Риск технической компании:       {data.get('technical', {}).get('conclusion', '—')}",
        ]
        return lines

    def _general(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        eg = agent.get("agent_egrul", {})
        return [
            f"Наименование:          {eg.get('full_name', '—') or eg.get('short_name', '—')}",
            f"ИНН:                   {self.inn}",
            f"ОГРН:                  {eg.get('ogrn', '—')}",
            f"КПП:                   {eg.get('kpp', '—')}",
            f"Дата регистрации:      {eg.get('reg_date', '—')}",
            f"Налоговый режим:       {eg.get('tax_regime', '—')}",
            f"Юридический адрес:     {eg.get('address', '—')}",
            f"ОКВЭД:                 {eg.get('okved_main', '—')}",
            f"Статус:                {eg.get('status', '—')}",
            f"Уставной капитал:      {eg.get('authorized_capital', '—')} ₽",
            f"Сотрудников:           {eg.get('employees', '—')}",
        ]

    def _management(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        eg = agent.get("agent_egrul", {})
        director = eg.get("director", {})
        disqual = data.get("disqual_data", {})
        lines = [
            f"Директор:              {director.get('fio', '—')}",
            f"ИНН директора:         {director.get('inn', '—')}",
            f"Дисквалификация:       {'да' if disqual.get('has_disqualified') else 'нет'}",
            "",
            "Учредители (участники):",
        ]
        founders = eg.get("founders", [])
        if not founders:
            lines.append("  — данные не найдены в ЕГРЮЛ")
        else:
            for f in founders:
                lines.append(f"  • {f.get('fio', '—')} (ИНН: {f.get('inn', '—')}) — доля: {f.get('share', '—')}")
        if disqual.get("disqualified"):
            lines.append("")
            lines.append("Дисквалифицированные лица:")
            for d in disqual["disqualified"]:
                lines.append(f"  • {d.get('fio', '—')}: {d.get('reason', '—')}")
        return lines

    def _related(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        cross = agent.get("agent_crosscheck", {})
        conns = cross.get("connections", [])
        if not conns:
            return ["Связанные компании не выявлены."]
        return [f"  • {c.get('type')}: {c.get('detail')} [{c.get('strength')}]" for c in conns[:10]]

    def _finance(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        tp = agent.get("agent_transparent", {})
        rp = agent.get("agent_rusprofile", {})
        eg = agent.get("agent_egrul", {})
        rev = rp.get("revenue", {})
        profit = rp.get("net_profit", {})

        rev_2022 = rev.get("2022", 0)
        rev_2023 = rev.get("2023", 0)
        rev_2024 = rev.get("2024", 0)
        pr_2022 = profit.get("2022", 0)
        pr_2023 = profit.get("2023", 0)
        pr_2024 = profit.get("2024", 0)

        employees = tp.get("employees_count", eg.get("employees", "—"))
        tax_burden = tp.get("tax_burden_percent", eg.get("tax_burden_percent", "—"))
        eg_revenue = eg.get("revenue_2025", "")
        eg_profit = eg.get("profit_2025", "")

        lines = ["Финансовые показатели:"]
        if rev_2022 or rev_2023 or rev_2024:
            lines += [
                f"  Выручка 2022:  {self._fmt(rev_2022)}",
                f"  Выручка 2023:  {self._fmt(rev_2023)}",
                f"  Выручка 2024:  {self._fmt(rev_2024)}",
                f"  Прибыль 2022:  {self._fmt(pr_2022)}",
                f"  Прибыль 2023:  {self._fmt(pr_2023)}",
                f"  Прибыль 2024:  {self._fmt(pr_2024)}",
            ]
        else:
            lines += [
                f"  Выручка 2025:  {eg_revenue if eg_revenue else '—'}",
                f"  Прибыль 2025:  {eg_profit if eg_profit else '—'}",
            ]
        lines += [
            f"  Чистые активы: {'—'}",
            f"  Сотрудники:    {employees}",
            f"  Налоговая нагрузка: {tax_burden}%" if tax_burden != "—" else "  Налоговая нагрузка: —",
            f"  Налоги уплачено: {self._fmt(eg.get('tax_paid', '—'))}",
            f"  Налоговая задолженность: {self._fmt(eg.get('tax_debt', '—'))}",
        ]
        return lines

    def _court(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        ar = agent.get("agent_arbitr", {})
        cases = ar.get("cases", [])
        if not cases:
            return ["Судебные дела не найдены."]
        return [
            f"  • №{c.get('number', '—')} от {c.get('date', '—')} — сумма: {self._fmt(c.get('amount', 0))} (роль: {c.get('role', '—')})"
            for c in cases[:10]
        ]

    def _enforcement(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        fssp = agent.get("agent_fssp", {})
        procs = fssp.get("proceedings", [])
        if not procs:
            return ["Исполнительные производства не найдены."]
        return [
            f"  • №{p.get('number', '—')} — долг: {self._fmt(p.get('amount', 0))} — статус: {p.get('status', '—')}"
            for p in procs[:10]
        ]

    def _licenses(self, data: dict) -> list:
        lic = data.get("licenses_data", {})
        if not lic.get("license_required"):
            return ["Данные о лицензиях отсутствуют."]
        lines = [f"Требуется лицензия: да ({lic.get('license_activity', '—')})"]
        if lic.get("has_license"):
            for l in lic.get("licenses", []):
                lines.append(f"  • {l.get('number', '—')} — {l.get('activity', '—')}")
        else:
            lines.append("  — лицензия не найдена —")
        return lines

    def _add_risks_table(self, data: dict):
        p = self.doc.add_paragraph()
        run = p.add_run("\n9. Оценка рисков")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 51, 102)

        sc = data.get("scoring", {})
        table = self.doc.add_table(rows=5, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(["Вид риска", "Оценка", "Уровень"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for idx, (name, score_key, level_key) in enumerate([
            ("Налоговый", "tax_score", "tax_level"),
            ("Финансовый", "financial_score", "financial_level"),
            ("Корпоративный", "corporate_score", "corporate_level"),
            ("Репутационный", "reputation_score", "reputation_level"),
        ]):
            row = table.rows[idx + 1]
            row.cells[0].text = name
            row.cells[1].text = f"{sc.get(score_key, 0)}/100"
            row.cells[2].text = sc.get(level_key, "—")

        splitting = data.get("splitting", {})
        technical = data.get("technical", {})
        self.doc.add_paragraph(f"Риск дробления бизнеса: {splitting.get('conclusion', '—')}")
        self.doc.add_paragraph(f"Риск технической компании: {technical.get('conclusion', '—')}")

    def _fns_claims(self, data: dict) -> list:
        sc = data.get("scoring", {})
        splitting = data.get("splitting", {})
        technical = data.get("technical", {})
        lines = ["На основании выявленных факторов:"]
        if sc.get("tax_score", 0) >= 50:
            lines.append("  • Высокий налоговый риск — возможны претензии по необоснованной налоговой выгоде")
        if sc.get("corporate_score", 0) >= 50:
            lines.append("  • Высокий корпоративный риск — возможна инициация налоговой проверки")
        if splitting.get("conclusion") == "ВЫСОКАЯ":
            lines.append("  • Высокая вероятность дробления бизнеса — возможны претензии по ст. 54.1 НК РФ")
        if technical.get("conclusion") == "ВЫСОКАЯ":
            lines.append("  • Компания имеет признаки технической — возможно доначисление налогов")
        if len(lines) == 1:
            lines.append("  • Существенных предпосылок для претензий ФНС не выявлено.")
        return lines

    def _recommendations(self, data: dict) -> list:
        recs = data.get("recommendations", [])
        if recs:
            return [f"  • {r.replace('🔴 СРОЧНО: ', '').replace('🟡 ВАЖНО: ', '').replace('🟢 ЖЕЛАТЕЛЬНО: ', '')}" for r in recs]
        return ["  • Существенных рисков не выявлено. Рекомендуется регулярный мониторинг."]

    def _sources(self, data: dict) -> list:
        agent = data.get("agent_data", {})
        lines = []
        for name, src_data in agent.items():
            err = src_data.get("error")
            lines.append(f"  {'✓' if not err else '✗'} {src_data.get('source', name)} {'(недоступен)' if err else ''}")
        return lines

    def _save(self) -> str:
        reports_dir = os.path.join(BASE_DIR, REPORTS_DIR)
        os.makedirs(reports_dir, exist_ok=True)
        safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in self.company_name)[:40].strip()
        filename = f"{safe_name}_отчет_{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
        filepath = os.path.join(reports_dir, filename)
        self.doc.save(filepath)
        return filepath


def build(company_name: str, inn: str, data: dict) -> str:
    builder = ReportBuilder(company_name, inn)
    return builder.build(data)
